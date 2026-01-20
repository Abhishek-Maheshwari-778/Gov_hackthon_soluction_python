from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    document = Document()
    
    # Title Page
    title = document.add_heading('AADHAAR SEVA PULSE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = document.add_paragraph('Smart Resource Allocation & Anomaly Detection System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    document.add_paragraph('\n\n')
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Online Hackathon on Data-driven Innovation on Aadhaar 2026')
    run.bold = True
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Team ID: ').bold = True
    p.add_run('UIDAI_10441')
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Repository: ').bold = True
    p.add_run('https://github.com/Abhishek-Maheshwari-778/Gov_hackthon_soluction_python')
    
    document.add_page_break()
    
    # Exec Summary
    document.add_heading('1. Executive Summary', level=1)
    document.add_paragraph(
        '"Aadhaar Seva Pulse" is a pioneering analytics framework designed to shift government decision-making '
        'from Reactive to Predictive. By leveraging the high-velocity data of Aadhaar Demographics and Biometric '
        'updates, our solution creates a real-time "Pulse" of Indian society.'
    )
    document.add_paragraph(
        'We utilize Unsupervised Machine Learning (K-Means Clustering & Isolation Forest) to segment districts '
        'into strategic profiles and identify operational anomalies.'
    )
    
    # Problem Statement
    document.add_heading('2. Problem Statement', level=1)
    document.add_paragraph('1. Urban Overload: Rapidly developing "Magnet Districts" face huge migration inflows.')
    document.add_paragraph('2. The Digital Divide: Rural districts with high birth rates lag in mandatory biometric updates.')
    
    # Methodology (Images)
    document.add_heading('3. Methodology & Analysis', level=1)
    
    # Image 1
    document.add_heading('3.1 Data Density & Missing Values', level=2)
    img_path = r'graph and charts/Outlier_Check_Enrolments_and_Distribution_ChildEnrolments0_5.png'
    if os.path.exists(img_path):
        document.add_picture(img_path, width=Inches(6.0))
        document.add_paragraph('Figure 1: Data Density Heatmap', style='Caption')
    else:
        document.add_paragraph('[Image missing: ' + img_path + ']')

    # Image 2
    document.add_heading('3.2 Demographic Composition', level=2)
    img_path = r'graph and charts/Enrolment_Composition_by_Age.png'
    if os.path.exists(img_path):
        document.add_picture(img_path, width=Inches(6.0))
        document.add_paragraph('Figure 2: Age Breakdown', style='Caption')

    # Image 3
    document.add_heading('3.3 State Performance', level=2)
    img_path = r'graph and charts/Top_10_States_Migration_vs_Compliance.png'
    if os.path.exists(img_path):
        document.add_picture(img_path, width=Inches(6.0))
        document.add_paragraph('Figure 3: Migration vs Compliance per State', style='Caption')

    # Image 4
    document.add_heading('3.4 Correlation Analysis', level=2)
    img_path = r'graph and charts/Research_Correlation_Matrix.png'
    if os.path.exists(img_path):
        document.add_picture(img_path, width=Inches(5.0))
        document.add_paragraph('Figure 4: Correlation Matrix', style='Caption')

    # Image 5
    document.add_heading('3.5 District Clustering (K-Means)', level=2)
    img_path = r'graph and charts/District_Segmentation_(AI_Clustering).png'
    if os.path.exists(img_path):
        document.add_picture(img_path, width=Inches(6.0))
        document.add_paragraph('Figure 5: 4 Strategic Clusters of Districts', style='Caption')

    document.add_page_break()

    # Findings (Table)
    document.add_heading('4. Key Findings & Anomalies', level=1)
    document.add_paragraph('Our Isolation Forest algorithm detected 23 critical anomalies. The top 5 are listed below:')
    
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'State'
    hdr_cells[1].text = 'District'
    hdr_cells[2].text = 'Migration Volume'
    hdr_cells[3].text = 'New Births'
    
    # Data from findings
    data = [
        ('Andhra Pradesh', 'East Godavari', '136,935', '7,161'),
        ('Andhra Pradesh', 'Guntur', '155,195', '8,769'),
        ('Andhra Pradesh', 'Kurnool', '148,267', '10,923'),
        ('Andhra Pradesh', 'Visakhapatnam', '137,329', '7,688'),
        ('Bihar', 'Sitamarhi', '157,203', '20,679')
    ]
    
    for state, dist, mig, birth in data:
        row_cells = table.add_row().cells
        row_cells[0].text = state
        row_cells[1].text = dist
        row_cells[2].text = mig
        row_cells[3].text = birth
        
    document.add_paragraph('\nAction Item: Immediate audit of Registrars in these red-flagged zones.')

    # Conclusion
    document.add_heading('5. Conclusion', level=1)
    document.add_paragraph(
        'The tools developed in this project (Recursive Loaders, Cluster Models, Anomaly Detectors) '
        'are modular and scalable. They can be integrated into UIDAI\'s central command center to ensure '
        'that No Resident is Left Behind.'
    )
    
    filename = 'AADHAARSEVAPULSE_Team_UIDAI_10441.docx'
    document.save(filename)
    print(f"Successfully generated {filename}")

if __name__ == "__main__":
    create_report()
